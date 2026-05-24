from docx import Document
import os

file_path = "chap9_1141222.docx"

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
    exit(1)

try:
    doc = Document(file_path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    print("\n--- First 20 Paragraphs ---")
    for i, p in enumerate(doc.paragraphs[:20]):
        if p.text.strip():
            print(f"{i}: {p.text.strip()}")
            
    print("\n--- First 2 Tables ---")
    for i, table in enumerate(doc.tables[:2]):
        print(f"Table {i}:")
        for row in table.rows[:5]: # First 5 rows
            print([cell.text.strip() for cell in row.cells])
            
except Exception as e:
    print(f"Error reading docx: {e}")
